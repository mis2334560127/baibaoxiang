"""
PDF 转 Word 模块
支持两种模式：
1. 数字型 PDF（文字可选层）→ pdf2docx 直接转换
2. 扫描型 PDF（图片）→ PyMuPDF 渲染 + PaddleOCR 识别
"""
import os
import io
from pathlib import Path


def detect_pdf_type(file_path: str) -> tuple[str, int, int]:
    """
    检测 PDF 类型。

    Returns:
        (type_str, text_pages, total_pages)
        type_str: "text" / "scanned" / "mixed"
    """
    try:
        import fitz
    except ImportError:
        return ("text", 0, 0)

    try:
        doc = fitz.open(file_path)
        total = doc.page_count
        text_pages = 0
        for i in range(total):
            page = doc[i]
            text = page.get_text().strip()
            if text:  # 非空白文本即视为有文字层
                text_pages += 1
        doc.close()

        if text_pages == total:
            return ("text", text_pages, total)
        elif text_pages == 0:
            return ("scanned", 0, total)
        else:
            return ("mixed", text_pages, total)
    except Exception:
        return ("text", 0, 0)


def _pdf_page_has_tables(file_path: str) -> bool:
    """
    使用 PyMuPDF 检测 PDF 中是否包含表格（pdf2docx 表格还原不佳时启用 OCR 兜底）。

    检测策略：
    1. 标准表格：PyMuPDF page.find_tables()
    2. 复杂框线 / 斜线表头表格：find_tables() 对斜线表头、复杂框线识别弱，
       通过页面绘制线条数量 + 表格关键词（如"类别"、"部别"、"总计"）二次判定。
    """
    try:
        import fitz
        doc = fitz.open(file_path)
        for page in doc:
            # 1. 标准表格检测
            try:
                tables = page.find_tables()
                if tables and tables.tables:
                    doc.close()
                    return True
            except Exception:
                # 旧版 PyMuPDF 可能没有 find_tables，忽略
                pass

            # 2. 斜线表头 / 复杂框线表格检测
            try:
                text = page.get_text()
                if len(text.strip()) > 20:
                    # 2a. 斜线表头关键词：类别+部别，或同时出现多个总计/合计
                    has_header_keywords = (
                        ("类别" in text and "部别" in text)
                        or ("项目" in text and "内容" in text)
                    )
                    # 2b. 绘制线条数量：普通表格至少 6-8 根线，斜线表头更多
                    drawings = page.get_drawings()
                    line_count = 0
                    for d in drawings:
                        items = d.get("items") if isinstance(d, dict) else None
                        if items:
                            for it in items:
                                it_type = (
                                    it[0]
                                    if isinstance(it, (list, tuple)) and len(it) > 0
                                    else getattr(it, "type", None)
                                )
                                if it_type in ("l", "line"):
                                    line_count += 1
                        elif isinstance(d, dict) and d.get("type") in ("l", "line"):
                            line_count += 1

                    # 关键词或大量线条都视为含复杂表格
                    if has_header_keywords or line_count >= 8:
                        doc.close()
                        return True
            except Exception:
                pass
        doc.close()
    except Exception:
        pass
    return False



def convert_pdf_to_docx(
    file_path: str,
    output_dir: str = "",
    preserve_formatting: bool = True,
    preserve_images: bool = True,
    force_ocr: bool = False,
    ocr_lang: str = "ch",
    pure_ocr: bool = False,
    progress_callback: callable = None,
) -> str:
    """
    将 PDF 转换为 Word (.docx) 文档。
    自动检测 PDF 类型，扫描型使用 PaddleOCR（保留排版），文字型使用 pdf2docx。

    Args:
        file_path:             PDF 文件路径
        output_dir:            输出目录（空则同源目录）
        preserve_formatting:   保留原始格式（仅 pdf2docx 模式生效）
        preserve_images:       保留图片（仅 pdf2docx 模式生效）
        force_ocr:             强制使用 OCR（即使检测为文字型）
        ocr_lang:              OCR 识别语言（默认 "ch"，可选 "en" 等）
        pure_ocr:              纯 OCR 文字识别（不处理排版，直接 output 纯文本）
        progress_callback:     页级进度回调 callable(current_page, total_pages)，可为 None

    Returns:
        输出 .docx 文件路径
    """
    if not file_path.lower().endswith('.pdf'):
        raise ValueError(f"不支持的文件格式: {file_path}")

    if not os.path.exists(file_path):
        raise FileNotFoundError(f"文件不存在: {file_path}")

    # 确定输出路径（带可写性检测 + 自动降级）
    base = os.path.splitext(os.path.basename(file_path))[0]
    out_dir = output_dir or os.path.dirname(file_path)
    out_dir = _ensure_writable_dir(out_dir)
    out_path = _resolve_unique_output_path(out_dir, base)

    # 检测 PDF 类型
    pdf_type, text_pages, total_pages = detect_pdf_type(file_path)

    if pure_ocr:
        # 纯 OCR 模式：逐页识别文字，不做排版处理
        return _convert_with_pure_ocr(
            file_path, out_path, progress_callback=progress_callback, lang=ocr_lang,
        )
    elif not force_ocr and pdf_type == "text":
        # 文字型 PDF 如果包含表格，pdf2docx 经常把表格拆成文字。
        # 优先尝试 OCR 路径（已有表格检测/重建），若 PaddleOCR 未安装则回退 pdf2docx。
        if _pdf_page_has_tables(file_path):
            try:
                return _convert_with_ocr(
                    file_path, out_path,
                    pdf_type=pdf_type, text_pages=text_pages,
                    lang=ocr_lang, progress_callback=progress_callback,
                )
            except _PaddleOCRError:
                pass  # 没有 PaddleOCR，回退到 pdf2docx
        # 文字型：使用 pdf2docx 直接转换
        return _convert_with_pdf2docx(
            file_path, out_path, out_dir,
            preserve_formatting, preserve_images,
            progress_callback=progress_callback,
        )
    else:
        # 扫描型 / 混合型 / 强制 OCR：使用 OCR 转换（保留排版）
        return _convert_with_ocr(
            file_path, out_path,
            pdf_type=pdf_type, text_pages=text_pages,
            lang=ocr_lang, progress_callback=progress_callback,
        )


def _convert_with_pdf2docx(
    file_path: str, out_path: str, out_dir: str,
    preserve_formatting: bool, preserve_images: bool,
    progress_callback: callable = None,
) -> str:
    """使用 pdf2docx 库转换文字型 PDF（含后处理修正）"""
    def _report(page: int, total: int):
        if progress_callback:
            try:
                progress_callback(page, total)
            except Exception:
                pass

    try:
        from pdf2docx import Converter

        images_folder = None
        if preserve_images:
            images_folder = os.path.join(out_dir, "images")
            os.makedirs(images_folder, exist_ok=True)

        cv = Converter(file_path)
        _report(0, 1)  # 开始（pdf2docx 是黑盒，只能报告开始/结束）
        cv.convert(
            out_path,
            start=0,
            end=None,
            multi_processing=False,
            image_folder=images_folder,
        )
        cv.close()
    except ImportError:
        raise ImportError("pdf2docx 库未安装。请运行: pip install pdf2docx")
    except Exception as e:
        raise RuntimeError(f"PDF 转换失败: {str(e)}")

    if not os.path.exists(out_path):
        raise RuntimeError("转换似乎失败，输出文件未生成")

    # ── 后处理：根据用户选项决定行为 ──
    if preserve_formatting:
        # 保留原始格式：仅做轻量级关键修复（边距 + 极端字号保护），不破坏字体/颜色/粗体等
        _light_post_process_docx(out_path)
    else:
        # 不保留格式：标准化排版（统一宋体、清除加粗/颜色、规范行距）
        _post_process_docx_formatting(out_path)

    _report(1, 1)  # 完成
    return out_path


def _light_post_process_docx(docx_path: str):
    """
    轻量级后处理（勾选「保留原始格式」时使用）。

    只修关键问题，不破坏 PDF 原始字体/颜色/粗体等排版信息：
    1. 页面边距：保留 pdf2docx 设置的原始边距（不修改）
    2. 极端字号保护：< 5pt → 8pt，> 30pt → 24pt（防止不可读或溢出）
    3. 中文正文首行缩进：对长段落自动检测并补 2 字符缩进
    4. 段落间距标准化：消除异常大的段前/段后间距
    """
    try:
        from docx import Document
        from docx.shared import Pt, Emu
        from docx.oxml.ns import qn
    except ImportError:
        return

    try:
        doc = Document(docx_path)

        # ══════════════════════════════════════
        # 第1步：极端字号保护（不修改字体/颜色/粗体/对齐/缩进）
        # ══════════════════════════════════════
        for para in doc.paragraphs:
            for run in para.runs:
                if run.font.size:
                    fs_pt = run.font.size.pt
                    if fs_pt < 5:
                        run.font.size = Pt(8)
                    elif fs_pt > 30:
                        run.font.size = Pt(24)

        # ══════════════════════════════════════
        # 第2步：计算全文"正文字号"参考值（用于缩进量计算）
        # ══════════════════════════════════════
        body_font_sizes: list[float] = []
        for para in doc.paragraphs:
            text = para.text.strip()
            if not text:
                continue
            cjk_count = sum(
                1 for ch in text
                if '\u4e00' <= ch <= '\u9fff' or '\u3400' <= ch <= '\u4dbf'
            )
            if cjk_count < 5:  # 太短的跳过（标题/编号行）
                continue
            # 收集该段落首 run 的字号
            for run in para.runs:
                if run.font.size and run.text.strip():
                    body_font_sizes.append(run.font.size.pt)
                    break

        # 中位数作为正文字号，无数据时默认 10.5pt
        if body_font_sizes:
            body_font_sizes.sort()
            body_fs = body_font_sizes[len(body_font_sizes) // 2]
        else:
            body_fs = 10.5

        # ══════════════════════════════════════
        # 第3步：中文正文首行缩进（2字符）
        # ══════════════════════════════════════
        indent_pt = body_fs * 2.0  # 2个中文字符宽度

        for para in doc.paragraphs:
            text = para.text.strip()
            if not text:
                continue

            # 3a. 跳过标题样式
            style_name = (para.style.name if para.style else "").lower()
            if "heading" in style_name or "标题" in style_name:
                continue

            # 3b. 跳过居中/右对齐（标题特征）
            pf = para.paragraph_format
            try:
                from docx.enum.text import WD_ALIGN_PARAGRAPH
                if pf.alignment in (WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.RIGHT):
                    continue
            except Exception:
                pass

            # 3c. 判断是否为中文正文段落（CJK占比 > 40% 且长度 ≥ 20字）
            #     同时跳过列表项，避免把“1. / A. / 材料一”也加上首行缩进
            cjk_count = sum(
                1 for ch in text
                if '\u4e00' <= ch <= '\u9fff' or '\u3400' <= ch <= '\u4dbf'
            )
            total_chars = len(text)
            if total_chars < 20 or cjk_count / max(total_chars, 1) < 0.4:
                continue
            if _is_list_item(text):
                continue

            # 3d. 已有缩进则跳过（保护 pdf2docx 已正确设置的缩进）
            existing = pf.first_line_indent
            if existing is not None:
                existing_pt = existing.pt if hasattr(existing, 'pt') else existing / 12700 * 72
                if existing_pt > 2.0:  # 已有 ≥ 2pt 的正缩进
                    continue

            # 3e. 左缩进检测：如果左缩进 > 1字符宽，说明 pdf2docx 用左缩进代替了首行缩进
            left_indent = pf.left_indent
            if left_indent is not None:
                left_pt = left_indent.pt if hasattr(left_indent, 'pt') else left_indent / 12700 * 72
                if left_pt > body_fs * 1.5:
                    # 把左缩进转换为首行缩进（更符合中文排版习惯）
                    pf.left_indent = Pt(0)
                    pf.first_line_indent = Pt(indent_pt)
                    continue

            # 3f. 设置首行缩进
            pf.first_line_indent = Pt(indent_pt)

        # ══════════════════════════════════════
        # 第4步：段落间距标准化
        # ══════════════════════════════════════
        for para in doc.paragraphs:
            pf = para.paragraph_format
            if pf.space_before is not None and pf.space_before > Pt(18):
                pf.space_before = Pt(6)
            if pf.space_after is not None and pf.space_after > Pt(18):
                pf.space_after = Pt(6)
            # 负间距也修正
            if pf.space_before is not None and pf.space_before < Pt(0):
                pf.space_before = Pt(0)
            if pf.space_after is not None and pf.space_after < Pt(0):
                pf.space_after = Pt(0)

        doc.save(docx_path)
    except Exception:
        pass


def _post_process_docx_formatting(docx_path: str):
    """
    后处理 pdf2docx 输出的 DOCX（未勾选「保留原始格式」时使用）。

    标准化排版成规范的中文文档格式：
    1. 页面边距：A4 标准 2.54cm 四边
    2. 正文：SimSun 10.5pt，首行缩进 2 字符，1.5 倍行距
    3. 标题：保留粗体，SimHei（黑体）风格，清除奇怪颜色
    4. 清除全文不一致的加粗和颜色（除标题外）
    5. 段落间距标准化
    """
    try:
        from docx import Document
        from docx.shared import Pt, Cm
        from docx.oxml.ns import qn
        from docx.enum.text import WD_ALIGN_PARAGRAPH
    except ImportError:
        return  # 后处理非必需，静默跳过

    try:
        doc = Document(docx_path)

        # ══════════════════════════════════════
        # 第1步：A4 标准页面边距
        # ══════════════════════════════════════
        for section in doc.sections:
            section.top_margin = Cm(2.54)
            section.bottom_margin = Cm(2.54)
            section.left_margin = Cm(2.54)
            section.right_margin = Cm(2.54)

        INDENT_2CHAR = Pt(21.0)   # 2 字符 = 2 × 10.5pt
        BODY_FS = Pt(10.5)        # 五号 = 10.5pt

        # ══════════════════════════════════════
        # 第2步：遍历段落，分类处理
        # ══════════════════════════════════════
        for para in doc.paragraphs:
            style_name = (para.style.name if para.style else "").lower()
            is_heading = ("heading" in style_name or "标题" in style_name)

            text = para.text.strip()
            # 判断段落正文特征：CJK 长文本 + 非标题
            cjk_count = sum(
                1 for ch in text
                if '\u4e00' <= ch <= '\u9fff' or '\u3400' <= ch <= '\u4dbf'
            )
            is_body_like = (
                not is_heading
                and len(text) >= 15
                and (cjk_count / max(len(text), 1)) > 0.3
                and not _is_list_item(text)
            )

            # ── 处理每个 run ──
            for run in para.runs:
                # --- 字号保护 ---
                if run.font.size:
                    fs_pt = run.font.size.pt
                    if is_heading:
                        if fs_pt < 9:
                            run.font.size = BODY_FS
                        elif fs_pt > 26:
                            run.font.size = Pt(22)
                    else:
                        if fs_pt < 6 or fs_pt > 18:
                            run.font.size = BODY_FS

                # --- 统一字体 ---
                if is_heading:
                    run.font.name = 'SimHei'
                    _set_east_asian_font(run, 'SimHei')
                else:
                    run.font.name = 'SimSun'
                    _set_east_asian_font(run, 'SimSun')

                # --- 加粗：标题保留，正文清除 ---
                if not is_heading:
                    run.font.bold = None

                # --- 清除文字颜色 ---
                try:
                    rpr = run._element.get_or_add_rPr()
                    for color_tag in [qn('w:color'), qn('w14:color')]:
                        el = rpr.find(color_tag)
                        if el is not None:
                            rpr.remove(el)
                except Exception:
                    pass

                # --- 清除无效字体回退 ---
                try:
                    rpr = run._element.get_or_add_rPr()
                    rFonts = rpr.find(qn('w:rFonts'))
                    if rFonts is not None:
                        for attr in ['w:ascii', 'w:hAnsi', 'w:cs']:
                            rFonts.attrib.pop(qn(attr), None)
                except Exception:
                    pass

            # ── 段落格式 ──
            pf = para.paragraph_format

            # 行间距
            if not is_heading:
                try:
                    pf.line_spacing = 1.5
                except Exception:
                    pass

            # 首行缩进：只对正文段落
            if is_body_like:
                try:
                    existing = pf.first_line_indent
                    if existing is None or (
                        hasattr(existing, 'pt') and existing.pt < 2.0
                    ):
                        pf.first_line_indent = INDENT_2CHAR
                except Exception:
                    pass

            # 清除多余的左侧缩进（正文段落不应该有大左缩进）
            if is_body_like and pf.left_indent is not None:
                try:
                    left_pt = pf.left_indent.pt if hasattr(pf.left_indent, 'pt') else pf.left_indent / 12700 * 72
                    if left_pt > 5:
                        pf.left_indent = Pt(0)
                except Exception:
                    pass

            # 段落间距标准化
            try:
                if pf.space_before is not None and pf.space_before > Pt(12):
                    pf.space_before = Pt(6)
                if pf.space_after is not None and pf.space_after > Pt(12):
                    pf.space_after = Pt(6)
                if pf.space_before is not None and pf.space_before < Pt(0):
                    pf.space_before = Pt(0)
                if pf.space_after is not None and pf.space_after < Pt(0):
                    pf.space_after = Pt(0)
            except Exception:
                pass

        doc.save(docx_path)
    except Exception:
        pass  # 后处理失败不影响主流程


def _ensure_writable_dir(dir_path: str) -> str:
    """
    确保目录可写。不可写时自动降级到安全目录（桌面/百宝箱输出）。

    降级场景：
    - 源 PDF 在只读/系统目录（如 D:/视频/）
    - output_dir 未配置且源目录无写权限
    - UAC 保护的特殊文件夹
    """
    # 1. 先尝试创建目标目录
    try:
        os.makedirs(dir_path, exist_ok=True)
        # 写入测试：尝试创建+删除一个临时文件
        test_file = os.path.join(dir_path, f".write_test_{os.getpid()}")
        with open(test_file, "w") as f:
            f.write("test")
        os.remove(test_file)
        return dir_path  # 可写，原样返回
    except (OSError, PermissionError):
        pass

    # 2. 降级到用户桌面/百宝箱输出
    fallback_candidates = [
        str(Path.home() / "Desktop" / "百宝箱输出"),
        str(Path.home() / "Documents" / "百宝箱输出"),
        str(Path.home() / "Downloads" / "百宝箱输出"),
        os.path.join(os.path.dirname(__file__), "..", "data", "output"),
    ]

    for candidate in fallback_candidates:
        try:
            os.makedirs(candidate, exist_ok=True)
            test_file = os.path.join(candidate, f".write_test_{os.getpid()}")
            with open(test_file, "w") as f:
                f.write("test")
            os.remove(test_file)
            print(f"[PDF] 输出目录不可写，已自动切换为: {candidate}")
            return candidate
        except (OSError, PermissionError):
            continue

    # 3. 最终兜底：临时目录（一定可写）
    import tempfile
    fallback = os.path.join(tempfile.gettempdir(), "百宝箱输出")
    os.makedirs(fallback, exist_ok=True)
    print(f"[PDF] 输出目录不可写，已回退到临时目录: {fallback}")
    return fallback


def _resolve_unique_output_path(out_dir: str, base: str) -> str:
    """
    生成一个可写入的输出路径。
    当目标文件被其他程序（如 Word / WPS）占用时，自动追加序号/时间戳避免覆盖。
    """
    ext = ".docx"
    candidate = os.path.join(out_dir, f"{base}{ext}")

    # 如果文件不存在，直接返回
    if not os.path.exists(candidate):
        return candidate

    # 尝试以写模式打开原文件，测试是否被占用
    try:
        with open(candidate, "a+b"):
            pass
        return candidate
    except (PermissionError, OSError):
        pass

    # 文件被占用：生成带时间戳的新文件名
    from datetime import datetime
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    candidate = os.path.join(out_dir, f"{base}_{timestamp}{ext}")
    if not os.path.exists(candidate):
        return candidate

    # 极端情况：连时间戳文件也存在，追加序号
    counter = 1
    while True:
        candidate = os.path.join(out_dir, f"{base}_{timestamp}_{counter}{ext}")
        if not os.path.exists(candidate):
            return candidate
        counter += 1


_PADDLEOCR_INSTANCE = None
_PADDLEOCR_LANG = None


class _PaddleOCRError(RuntimeError):
    """PaddleOCR 未安装或初始化失败的异常"""


def _init_paddleocr(lang: str = "ch"):
    """初始化 PaddleOCR（全局单例，按语言缓存），兼容不同版本参数差异"""
    global _PADDLEOCR_INSTANCE, _PADDLEOCR_LANG

    if _PADDLEOCR_INSTANCE is not None and _PADDLEOCR_LANG == lang:
        return _PADDLEOCR_INSTANCE

    try:
        import inspect
        from paddleocr import PaddleOCR
    except ImportError:
        raise _PaddleOCRError(
            "PaddleOCR 未安装。请运行：\n"
            "pip install paddlepaddle paddleocr\n"
            "首次使用会自动下载中文识别模型（约 100MB），请保持网络连接。"
        )


    # 根据当前 PaddleOCR 版本支持的参数动态构建 kwargs
    sig = inspect.signature(PaddleOCR.__init__)
    supported = set(sig.parameters.keys())

    kwargs: dict[str, object] = {"lang": lang}
    if "show_log" in supported:
        kwargs["show_log"] = False
    if "use_angle_cls" in supported:
        kwargs["use_angle_cls"] = True
    if "use_gpu" in supported:
        kwargs["use_gpu"] = False

    try:
        _PADDLEOCR_INSTANCE = PaddleOCR(**kwargs)
        _PADDLEOCR_LANG = lang
        return _PADDLEOCR_INSTANCE
    except Exception as e:
        raise _PaddleOCRError(f"PaddleOCR 初始化失败: {e}")


def _parse_paddleocr_result(result) -> list[tuple[str, list, float]]:
    """
    统一解析 PaddleOCR 不同版本的返回格式。

    支持格式：
    1. PaddleOCR 2.x: [[ [box], (text, conf) ], ...]
    2. PaddleOCR 2.x 变体: [[ [box], text, conf ], ...]
    3. PaddleOCR 3.x batch: [{ "dt_polys": [...], "rec_text": [...], "rec_score": [...] }, ...]
    4. PaddleOCR 3.x single: { "dt_polys": ..., "rec_text": ..., "rec_score": ... }

    Returns:
        [(text, bbox, conf), ...]
    """
    lines: list[tuple[str, list, float]] = []

    if not result:
        return lines

    # 3.x 单结果字典或对象
    if isinstance(result, dict) or hasattr(result, "dt_polys"):
        result = [result]

    # 3.x batch：每个元素是字典或类似对象
    if result and (isinstance(result[0], dict) or hasattr(result[0], "dt_polys")):
        for item in result:
            polys = item.get("dt_polys", []) if isinstance(item, dict) else getattr(item, "dt_polys", [])
            texts = item.get("rec_text", []) if isinstance(item, dict) else getattr(item, "rec_text", [])
            scores = item.get("rec_score", []) if isinstance(item, dict) else getattr(item, "rec_score", [])
            for poly, text, score in zip(polys, texts, scores):
                if text is None:
                    continue
                lines.append((str(text), poly, float(score) if isinstance(score, (int, float)) else 0.9))
        return lines

    # 2.x 格式
    page_result = result[0] if result else None
    if page_result is None:
        return lines

    for line_info in page_result:
        if not isinstance(line_info, (list, tuple)) or len(line_info) < 2:
            continue

        bbox = line_info[0]
        rec_info = line_info[1]

        if isinstance(rec_info, (list, tuple)) and len(rec_info) >= 2:
            text = rec_info[0]
            conf = rec_info[1]
        else:
            # 变体：[bbox, text, conf]
            if len(line_info) >= 3:
                text = line_info[1]
                conf = line_info[2]
            else:
                continue

        if text is None:
            continue
        lines.append((str(text), bbox, float(conf) if isinstance(conf, (int, float)) else 0.9))

    return lines


def _run_ocr(ocr, img_array):
    """
    调用 PaddleOCR 进行识别，兼容 2.x（ocr）和 3.x（predict）。
    PaddleOCR 3.x 已弃用 ocr()，优先使用 predict()。
    """
    if hasattr(ocr, "predict"):
        return ocr.predict(img_array)
    return ocr.ocr(img_array)


def check_paddleocr_available() -> tuple[bool, str]:
    """
    检查 PaddleOCR 是否可用。

    Returns:
        (available, message)
    """
    try:
        _init_paddleocr("ch")
        return True, "PaddleOCR 中文识别已就绪 ✓"
    except _PaddleOCRError as e:
        return False, str(e)
    except Exception as e:
        return False, f"PaddleOCR 不可用: {e}"


def get_ocr_languages() -> list[str]:
    """获取可用的 OCR 语言列表（PaddleOCR 支持的语言代码）"""
    return ["ch", "en", "chinese_cht", "japan", "korean", "french", "german"]


def _convert_with_pure_ocr(
    file_path: str, out_path: str,
    dpi: int = 400, lang: str = "ch",
    progress_callback: callable = None,
) -> str:
    """
    纯 OCR 文字识别模式：逐页渲染 PDF → PaddleOCR 识别 → 写入 Word。
    不进行任何排版分析（无分栏、无表格、无段落/标题分类），速度更快。
    """
    def _report(page: int, total: int):
        if progress_callback:
            try:
                progress_callback(page, total)
            except Exception:
                pass

    try:
        import fitz
        from PIL import Image
        from docx import Document
        from docx.shared import Pt
    except ImportError as e:
        raise ImportError(f"纯 OCR 转换缺少依赖: {e}")

    ocr = _init_paddleocr(lang)

    doc = fitz.open(file_path)
    total = doc.page_count

    docx = Document()
    style = docx.styles['Normal']
    style.font.name = 'SimSun'
    style.font.size = Pt(10.5)
    style.element.rPr.rFonts.set(
        '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}eastAsia', 'SimSun'
    )

    for i in range(total):
        _report(i + 1, total)
        page = doc[i]

        # 渲染页面为图片
        pix = page.get_pixmap(dpi=dpi)
        img_data = pix.tobytes("png")
        img = Image.open(io.BytesIO(img_data))

        # 图片预处理
        img = _preprocess_for_ocr(img)

        # PaddleOCR 识别（3.x 需要三通道输入）
        import numpy as np
        img_array = np.array(img.convert('RGB'))
        try:
            result = _run_ocr(ocr, img_array)
        except Exception as e:
            raise RuntimeError(f"OCR 识别失败 (第{i+1}页): {str(e)}")

        parsed = _parse_paddleocr_result(result)
        if not parsed:
            p = docx.add_paragraph("[此页未识别到文字内容]")
            p.add_run("").italic = True
            if i < total - 1:
                docx.add_page_break()
            continue

        # 按行输出，每行一个 Word 段落
        for text, _bbox, conf in parsed:
            if text and conf > 0.5:
                docx.add_paragraph(text.strip())

        # 页间分隔
        if i < total - 1:
            docx.add_page_break()

    doc.close()
    docx.save(out_path)

    if not os.path.exists(out_path):
        raise RuntimeError("输出文件未生成")

    return out_path


def _convert_with_ocr(
    file_path: str, out_path: str,
    pdf_type: str = "scanned", text_pages: int = 0,
    dpi: int = 400, lang: str = "ch",
    progress_callback: callable = None,
) -> str:
    """
    使用 PaddleOCR 转换扫描型 PDF。
    PaddleOCR 返回行级文字 + 坐标，用于重建段落排版。
    包含图片预处理（歪斜校正、自适应二值化）以提升识别准确率。
    """
    def _report(page: int, total: int):
        if progress_callback:
            try:
                progress_callback(page, total)
            except Exception:
                pass

    try:
        import fitz
        from PIL import Image
        from docx import Document
        from docx.shared import Pt
    except ImportError as e:
        raise ImportError(f"OCR 转换缺少依赖: {e}")

    import numpy as np

    ocr = _init_paddleocr(lang)

    doc = fitz.open(file_path)
    total = doc.page_count

    docx = Document()
    # 设置默认样式
    style = docx.styles['Normal']
    style.font.name = 'SimSun'
    style.font.size = Pt(10.5)
    style.element.rPr.rFonts.set(
        '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}eastAsia', 'SimSun'
    )

    for i in range(total):
        _report(i + 1, total)
        page = doc[i]

        # 混合型 PDF：有文字层的页直接提取文字
        page_text = page.get_text().strip()
        if pdf_type == "mixed" and page_text:
            docx.add_paragraph(page_text)
        else:
            # 渲染页面为高 DPI 图片
            pix = page.get_pixmap(dpi=dpi)
            img_data = pix.tobytes("png")
            img = Image.open(io.BytesIO(img_data))
            img_width, img_height = img.size

            # ---- 图片预处理：歪斜校正 + 自适应二值化 ----
            img = _preprocess_for_ocr(img, dpi=dpi)

            # ---- OCR 前版面分析：投影法检测栏目 ----
            col_count, col_boxes = _detect_page_columns_from_image(img)

            # PaddleOCR 对整页识别（3.x 需要三通道输入）
            img_array = np.array(img.convert('RGB'))
            try:
                result = _run_ocr(ocr, img_array)
            except Exception as e:
                raise RuntimeError(f"OCR 识别失败 (第{i+1}页): {str(e)}")

            parsed = _parse_paddleocr_result(result)
            if not parsed:
                p = docx.add_paragraph("[此页未识别到文字内容]")
                from docx.enum.text import WD_ALIGN_PARAGRAPH
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = p.add_run("")
                run.italic = True
                if i < total - 1:
                    docx.add_page_break()
                continue

            # ── 将 PaddleOCR 行级结果转为 line_info 格式 ──
            # 格式: (text, left, top, right, bottom, conf, line_index)
            ocr_lines: list[dict] = []
            for line_idx, (text, bbox, conf) in enumerate(parsed):
                if not text or not text.strip():
                    continue
                if conf < 0.5:
                    continue
                xs = [p[0] for p in bbox]
                ys = [p[1] for p in bbox]
                left = int(min(xs))
                top = int(min(ys))
                right = int(max(xs))
                bottom = int(max(ys))
                ocr_lines.append({
                    "text": text.strip(),
                    "left": left, "top": top,
                    "right": right, "bottom": bottom,
                    "width": right - left,
                    "height": bottom - top,
                    "conf": conf,
                    "line_idx": line_idx,
                })

            if not ocr_lines:
                if i < total - 1:
                    docx.add_page_break()
                continue

            # ── 全部文字行：原生行级管道重建排版 ──
            if col_count == 1:
                _build_from_ocr_lines(
                    docx, ocr_lines,
                    img_width, img_height, dpi,
                )
            else:
                for ci in range(col_count):
                    cx0, cx1 = col_boxes[ci]
                    col_width = cx1 - cx0
                    # 筛选属于该栏的行，并将坐标转为栏内相对坐标
                    col_lines = []
                    for ol in ocr_lines:
                        line_cx = ol["left"] + ol["width"] / 2
                        if cx0 <= line_cx <= cx1:
                            col_ol = dict(ol)
                            col_ol["left"] = ol["left"] - cx0
                            col_ol["right"] = ol["right"] - cx0
                            col_ol["width"] = ol["width"]
                            col_lines.append(col_ol)
                    if col_lines:
                        _build_from_ocr_lines(
                            docx, col_lines,
                            col_width, img_height, dpi,
                        )

        # 页间分隔
        if i < total - 1:
            docx.add_page_break()

    doc.close()
    docx.save(out_path)

    if not os.path.exists(out_path):
        raise RuntimeError("输出文件未生成")

    return out_path


def _deskew_image(img_array: "np.ndarray") -> "np.ndarray":
    """
    检测并校正扫描图片的歪斜角度。

    策略（WPS 级别）：
    1. Canny 边缘检测 → 霍夫概率直线检测
    2. 过滤：只取水平线（角度在 ±5° 内），排除短线（< img_width * 5%）
    3. 取角度中位数 → 仿射变换校正
    """
    import numpy as np
    try:
        import cv2
    except ImportError:
        return img_array

    h, w = img_array.shape[:2]
    h_float = float(h)
    w_float = float(w)

    # 灰度化
    if len(img_array.shape) == 3:
        gray = cv2.cvtColor(img_array, cv2.COLOR_BGR2GRAY)
    else:
        gray = img_array.copy()

    # Canny 边缘检测
    edges = cv2.Canny(gray, 50, 150, apertureSize=3)

    # 霍夫概率直线检测
    lines = cv2.HoughLinesP(edges, rho=1, theta=np.pi / 180,
                            threshold=int(w_float * 0.08),
                            minLineLength=int(w_float * 0.05),
                            maxLineGap=int(w_float * 0.02))
    if lines is None or len(lines) < 3:
        return img_array

    # 收集角度（兼容 OpenCV 不同版本的 lines shape）
    angles = []
    for line in lines:
        coords = line[0] if line.ndim > 1 and line.shape[0] == 1 else line
        x1, y1, x2, y2 = coords
        if abs(x2 - x1) < 5:
            continue  # 垂直线跳过
        angle = np.degrees(np.arctan2(y2 - y1, x2 - x1))
        if abs(angle) < 5:
            angles.append(angle)

    if len(angles) < 3:
        return img_array

    # 中位数角度
    median_angle = float(np.median(angles))
    if abs(median_angle) < 0.15:
        return img_array  # 很正，不需要校正

    # 仿射变换校正
    center = (w_float / 2, h_float / 2)
    matrix = cv2.getRotationMatrix2D(center, median_angle, 1.0)
    rotated = cv2.warpAffine(gray if len(img_array.shape) == 2 else img_array,
                              matrix, (w, h),
                              flags=cv2.INTER_CUBIC,
                              borderMode=cv2.BORDER_CONSTANT,
                              borderValue=(255, 255, 255))
    return rotated


def _adaptive_binarize(img_array: "np.ndarray") -> "np.ndarray":
    """
    自适应局部二值化（替代全局 Otsu）。

    对扫描件的关键优势：
    - 光照不均区域（页面边缘阴影）仍能正确二值化
    - 中文笔画不会被全局阈值一刀切
    - blockSize=31（约 0.5~1 个中文字大小）与 C=8 适合印刷体
    """
    import numpy as np
    try:
        import cv2
    except ImportError:
        # 回退到简单阈值
        _, binary = cv2.threshold(img_array, 0, 255, cv2.THRESH_BINARY + cv2.THRESH_OTSU)
        return binary

    # 先轻度高斯去噪
    denoised = cv2.GaussianBlur(img_array, (3, 3), 0)

    # 自适应高斯阈值
    binary = cv2.adaptiveThreshold(
        denoised, 255,
        cv2.ADAPTIVE_THRESH_GAUSSIAN_C,
        cv2.THRESH_BINARY,
        blockSize=31,
        C=8,
    )
    return binary


def _preprocess_for_ocr(img, enhance: bool = True, dpi: int = 400) -> "Image.Image":
    """
    WPS 级别的 OCR 预处理流水线。

    策略：
    1. 灰度化
    2. 歪斜检测 + 校正（霍夫直线法）
    3. 双边滤波去噪（保留边缘，去除扫描噪声）
    4. 自适应局部二值化（替代全局 Otsu，处理光照不均）
    5. 形态学膨胀（修复断裂笔画，不破坏字间距）
    6. 轻度锐化边缘

    尺寸不变，坐标不变。
    """
    import numpy as np
    try:
        import cv2
        HAS_CV2 = True
    except ImportError:
        HAS_CV2 = False

    from PIL import Image, ImageFilter, ImageOps, ImageMath

    if not enhance:
        if img.mode != 'L':
            img = img.convert('L')
        return img.convert('RGB')

    # 1. 转灰度
    if img.mode != 'L':
        img = img.convert('L')

    if not HAS_CV2:
        # 回退到旧版 PIL 纯处理
        img = img.filter(ImageFilter.GaussianBlur(radius=1))
        img = ImageOps.autocontrast(img, cutoff=1)
        return img

    # 2. PIL → numpy 数组
    arr = np.array(img, dtype=np.uint8)

    # 3. 歪斜检测 + 校正
    arr = _deskew_image(arr)

    # 4. 双边滤波去噪（保留边缘，比高斯模糊更适合文字）
    #    对于 400+ DPI 的高分辨率扫描件，使用更大的 kernel
    kernel_d = 5 if dpi <= 300 else 7
    arr = cv2.bilateralFilter(arr, kernel_d, 75, 75)

    # 5. 自适应局部二值化
    arr = _adaptive_binarize(arr)

    # 6. 形态学腐蚀（去除单独噪点）+ 膨胀（修复断裂笔画）
    #    对 400+ DPI，kernel 稍大一些
    kernel_size = 1 if dpi <= 300 else 2
    kernel = cv2.getStructuringElement(cv2.MORPH_ELLIPSE, (kernel_size, kernel_size))
    arr = cv2.morphologyEx(arr, cv2.MORPH_CLOSE, kernel)

    # 7. 轻度锐化（增强文字边缘对比度）
    #    kernel 3x3，权重 0.8
    sharpen_kernel = np.array([
        [-0.2, -0.2, -0.2],
        [-0.2,  2.6, -0.2],
        [-0.2, -0.2, -0.2]
    ], dtype=np.float32)
    arr = cv2.filter2D(arr, -1, sharpen_kernel)

    # 8. 钳制到 [0, 255] 并转回 PIL Image
    arr = np.clip(arr, 0, 255).astype(np.uint8)
    # PaddleOCR 3.x 需要三通道输入，灰度图转 RGB 不影响识别效果
    return Image.fromarray(arr, mode='L').convert('RGB')


def _detect_page_columns_from_image(img) -> tuple[int, list]:
    """
    基于图像垂直投影分析检测页面分栏（OCR 前）。

    使用预处理后的二值图片的垂直投影直方图，找到空白间隙来切分栏目。
    这比 OCR 后分析词坐标更准确，因为避免了 OCR 分割错误。

    Returns:
        (column_count, [(x0, x1), ...])  栏目数和每个栏目的 x 边界
    """
    import numpy as np
    try:
        import cv2
    except ImportError:
        img_width = img.width if hasattr(img, 'width') else img.shape[1]
        return 1, [(0, img_width)]

    # PIL Image → numpy 数组
    if hasattr(img, 'width'):
        arr = np.array(img, dtype=np.uint8)
    else:
        arr = img

    h, w = arr.shape[:2]
    if w < 200:
        return 1, [(0, w)]

    # 垂直投影：count non-zero (black) pixels in each column
    # 由于文字为黑色(0)背景为白色(255)，反转后 count non-zero
    if len(arr.shape) == 3:
        gray = cv2.cvtColor(arr, cv2.COLOR_BGR2GRAY)
    else:
        gray = arr

    # 反转：文字=255，背景=0
    binary = cv2.bitwise_not(gray) if np.mean(gray) > 127 else gray

    # 垂直投影
    v_proj = np.sum(binary, axis=0, dtype=np.float64)
    if v_proj.max() == 0:
        return 1, [(0, w)]

    # 归一化
    v_proj = v_proj / v_proj.max()

    # 忽略页面边缘（左右各 3% 边距）
    margin = int(w * 0.03)
    v_proj[:margin] = 0
    v_proj[-margin:] = 0

    # 找空白区（投影值 < 阈值的连续列）
    threshold = 0.02
    gap_margin = max(int(w * 0.02), 10)  # 最小间隙 2% 屏宽 或 10px
    gaps = []  # [(start, end)]
    in_gap = False
    gap_start = 0

    for x in range(w):
        if v_proj[x] < threshold:
            if not in_gap:
                gap_start = x
                in_gap = True
        else:
            if in_gap:
                gap_end = x
                if gap_end - gap_start >= gap_margin:
                    gaps.append((gap_start, gap_end))
                in_gap = False

    if in_gap:
        gap_end = w
        if gap_end - gap_start >= gap_margin:
            gaps.append((gap_start, gap_end))

    if not gaps:
        return 1, [(0, w)]

    # 只取页面中间区域的间隙（不能太靠左或太靠右）
    mid_gaps = [
        (s, e) for s, e in gaps
        if s > margin and e < w - margin
    ]

    if not mid_gaps:
        return 1, [(0, w)]

    # 按宽度排序，取最宽的间隙
    mid_gaps.sort(key=lambda g: g[1] - g[0], reverse=True)
    best_gap = mid_gaps[0]
    gap_center = (best_gap[0] + best_gap[1]) // 2

    # 确认左右两侧都有足够内容（至少 10% 页宽）才判定为双栏
    if gap_center > w * 0.15 and gap_center < w * 0.85:
        return 2, [
            (int(w * 0.02), gap_center),
            (gap_center, int(w * 0.98))
        ]

    return 1, [(0, w)]


# ──────────────────────────────────────────────
#  字号/粗体策略：扫描件 OCR 无法可靠推测字号和粗体
#  遵循 WPS 保守原则：全文统一字号、不加粗、不换字体、不换颜色
# ──────────────────────────────────────────────

# 扫描件默认正文字号（五号 ≈ 10.5pt，适用于中文文档）
_SCANNED_BODY_FS = 10.5





def _split_exam_options_in_line(text: str) -> list[str]:
    """
    把一行内同时出现多个 A/B/C/D 选项的文本拆成多行。
    例如："A. 旧石器时代早期 B. 旧石器时代晚期"
    → ["A. 旧石器时代早期", "B. 旧石器时代晚期"]
    """
    import re
    text = text.strip()
    if not text:
        return [""]

    # 匹配 A/B/C/D 后面跟标点或空格（不用 \b，因 re.split 中 lookahead+\b 行为异常）
    pattern = r"(?=[A-Da-d][\.．、])"
    parts = [p.strip() for p in re.split(pattern, text) if p.strip()]
    if len(parts) <= 1:
        return [text]

    # 如果第一个 part 不以选项开头，可能是题干，保留原样
    if not re.match(r"^[A-Da-d][\.．、]", parts[0]):
        return [text]

    return parts


def _split_multi_questions_in_line(text: str) -> list[str]:
    """
    把一行内多个题号的内容拆成多行。
    例如："25. 题干A  26. 题干B  27. 题干C"
    → ["25. 题干A", "26. 题干B", "27. 题干C"]

    同时处理括号题号：(1) xxx (2) yyy
    """
    import re
    text = text.strip()
    if not text:
        return [""]

    # 匹配题号模式：N.  N． (N)  或 材料X
    pattern = r"(?=\b\d{1,2}[\.．]\s+|\b\(\d{1,2}\)\s+|【答案】|【解析】|材料[一二三四五六七八九十]+)"
    parts = [p.strip() for p in re.split(pattern, text) if p.strip()]
    if len(parts) <= 1:
        return [text]

    # 过滤：如果第一个 part 不以题号开头，保留原样
    if not re.match(r"^\d{1,2}[\.．]|^\(\d{1,2}\)|^【|^材料", parts[0]):
        return [text]

    return parts


def _is_list_item(text: str) -> bool:
    """检测文本是否为列表项（编号、选项、材料等），避免被误加首行缩进。"""
    import re
    text = text.strip()
    if not text:
        return False

    # 编号项：1. 2. 3. 或 (1) (2) 或 ① 或 一、二、
    cjk_number = "一二三四五六七八九十"
    patterns = [
        r"^\d{1,2}[\.．、]\s*",                    # 1. 2. 3.
        r"^\(\d{1,2}\)\s*",                        # (1) (2)
        r"^[A-Da-d][\.．、]\s*",                    # A. B. C. D.
        r"^[①②③④⑤⑥⑦⑧⑨⑩]\s*",                    # ① ②
        rf"^[（(][{cjk_number}]+[）)]\s*",           # （一） (二)
        r"^【(?:答案|解析)】",                        # 【答案】【解析】
        r"^材料[一二三四五六七八九十]+\s*",          # 材料一 材料二
        r"^第[一二三四五六七八九十]+[章节条款]\s*",   # 第一章 第二节
    ]
    for pat in patterns:
        if re.match(pat, text):
            return True
    return False


def _detect_exam_boundary_line(text: str) -> bool:
    """
    检测文本行是否为考试卷常见边界（题号、选项、材料、答案等）。
    用于在 OCR 排版重建中强制分段，避免多题合并。
    """
    import re

    text = text.strip()
    if not text:
        return False

    patterns = [
        r"^\d{1,2}[\.．]\s*",             # 24.  25.
        r"^\(\d{1,2}\)\s*",               # (1)  (2)
        r"^【答案】",                       # 【答案】
        r"^【解析】",                       # 【解析】
        r"^材料[一二三四五六七八九十]+\s*",  # 材料一 材料二
        r"^[A-Da-d][\.．、]\s*",            # A.  B.
        r"^\(\d+分\)",                    # (12分) (15分)
    ]

    for pat in patterns:
        if re.match(pat, text):
            return True
    return False


def _build_column_from_ocr(
    docx, col_words: list, col_width: int, img_height: int, dpi: int,
):
    """
    从 Tesseract 词条数据重建单个栏目的 Word 排版。

    排版分析流水线：
    1. y坐标+重叠度 → 行分组（增强重叠判定）
    2. 逐行属性：字号（中英文差异）、左边界、粗体（密度法）、对齐
    3. 自适应密度聚类 → 段落边界
    4. 表格检测：x/y 网格对齐分析 → 原生 DOCX 表格
    5. 段落分类：位置权重 + 字号分布 + 编号模式 + 多特征融合
    6. 按分类输出（Heading1-3 / List / Body / Table）
    """
    img_width = col_width  # 复用原函数体内的变量名

    WordEntry = tuple[str, int, int, int, int, int, int]  # (text, left, top, width, height, block, line_num)

    from docx import Document
    from docx.shared import Pt, Cm, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn

    # ==================== 工具函数 ====================

    def _is_cjk(ch: str) -> bool:
        if not ch:
            return False
        cp = ord(ch)
        return (
            (0x4E00 <= cp <= 0x9FFF) or (0x3400 <= cp <= 0x4DBF) or
            (0x20000 <= cp <= 0x2A6DF) or (0x2F800 <= cp <= 0x2FA1F) or
            (0x3000 <= cp <= 0x303F) or (0xFF00 <= cp <= 0xFFEF) or
            (0x3040 <= cp <= 0x309F) or (0x30A0 <= cp <= 0x30FF) or
            (0xAC00 <= cp <= 0xD7AF) or (0xF900 <= cp <= 0xFAFF)
        )

    def _is_latin_or_digit(ch: str) -> bool:
        if not ch:
            return False
        cp = ord(ch)
        return (ord('a') <= cp <= ord('z') or
                ord('A') <= cp <= ord('Z') or
                ord('0') <= cp <= ord('9'))

    def _word_is_latin(w: str) -> bool:
        if not w:
            return False
        latin_count = sum(1 for ch in w if _is_latin_or_digit(ch))
        return latin_count > len(w) * 0.5

    def _join_words_smart(words: list[str]) -> str:
        if not words:
            return ""
        if len(words) == 1:
            return words[0]
        result = words[0]
        for i in range(1, len(words)):
            prev_latin = _word_is_latin(words[i - 1])
            curr_latin = _word_is_latin(words[i])
            if prev_latin and curr_latin:
                result += " " + words[i]
            else:
                result += words[i]
        return result

    raw_words = col_words
    if not raw_words:
        return


    # ==================== 第二步：按 y 坐标 + 增强重叠度 分组为行 ====================

    raw_words.sort(key=lambda w: w[2])  # 按 top 排序
    # 估算全局平均词高
    avg_word_height = sum(w[4] for w in raw_words) / len(raw_words)

    # 使用 Tesseract 的 block_num + line_num 辅助分组
    # 然后 y 重叠做修正
    lines_raw: list[list[WordEntry]] = []
    # 先按 (block_num, line_num) 粗分组
    from collections import OrderedDict
    line_map: dict[tuple[int, int], list[WordEntry]] = OrderedDict()
    for w in raw_words:
        key = (w[5], w[6])  # (block_num, line_num)
        if key not in line_map:
            line_map[key] = []
        line_map[key].append(w)

    # 把同一逻辑行的词合并，按 y 排序
    sorted_lines = sorted(line_map.values(),
                          key=lambda lst: min(w[2] for w in lst))

    # 二次合并：相邻逻辑行如果在 y 方向高度重叠 → 合并为一行
    merged_lines: list[list[WordEntry]] = []
    for line_words in sorted_lines:
        if not merged_lines:
            merged_lines.append(line_words)
            continue
        prev = merged_lines[-1]
        prev_ys = [w[2] for w in prev] + [w[2] + w[4] for w in prev]
        curr_ys = [w[2] for w in line_words] + [w[2] + w[4] for w in line_words]
        prev_mid = (min(prev_ys) + max(prev_ys)) / 2
        curr_mid = (min(curr_ys) + max(curr_ys)) / 2
        prev_range = max(prev_ys) - min(prev_ys)
        curr_range = max(curr_ys) - min(curr_ys)
        # 中位 y 差小于行高的 70% → 同一行（处理 Tesseract 将一行拆成多 block 的情况）
        if abs(curr_mid - prev_mid) < max(prev_range, curr_range, 1) * 0.7:
            merged_lines[-1].extend(line_words)
        else:
            merged_lines.append(line_words)

    # 每行内部按 left 排序
    for line in merged_lines:
        line.sort(key=lambda x: x[1])
    lines_raw = merged_lines

    # ==================== 第三步：计算每行的排版属性（统一字号，不推测粗体） ====================

    LineInfo = tuple[
        list[WordEntry],  # words (sorted by left)
        float,            # line_top (像素)
        float,            # line_bottom
        float,            # line_height
        float,            # font_size_pt (固定常量 _SCANNED_BODY_FS)
        float,            # left_margin (行左边界)
        float,            # right_margin (行右边界)
        bool,             # is_bold (固定 False，扫描件不推测)
        float,            # alignment_score (-1=left, 0=center, 1=right)
    ]

    line_infos: list[LineInfo] = []

    for line in lines_raw:
        # 行边界
        l_top = min(w[2] for w in line)
        l_bottom = max(w[2] + w[4] for w in line)
        l_height = l_bottom - l_top
        l_left = min(w[1] for w in line)
        l_right = max(w[1] + w[3] for w in line)

        # 不推测字号和粗体：扫描件不存在这些信息，统一使用固定常量
        fs = _SCANNED_BODY_FS

        # 对齐分数：左留白比例 vs 右留白比例
        left_space = l_left
        right_space = img_width - l_right
        total_space = left_space + right_space
        if total_space > img_width * 0.3:
            # 左右都有大留白 → 居中倾向
            balance = 1.0 - abs(left_space - right_space) / max(total_space, 1)
            align_score = 0.0 if balance > 0.6 else (-1.0 if left_space < right_space else 1.0)
        elif left_space < img_width * 0.05:
            align_score = -1.0  # 左对齐
        elif right_space < img_width * 0.05:
            align_score = 1.0   # 右对齐
        else:
            align_score = -1.0  # 默认左对齐

        line_infos.append((line, l_top, l_bottom, l_height, fs, l_left, l_right, False, align_score))

    # ==================== 第四步：自适应段落边界检测（密度聚类 + 多信号） ====================
    # 替代固定 1.5× 阈值。核心思想：相邻行的间距如果显著偏离其局部邻域的间距模式，
    # 则认为是段落边界。

    if len(line_infos) == 1:
        line_heights = [line_infos[0][3]]
        med_body_fs_single = line_infos[0][4]
        single_lines = [line_infos[0]]
        # 尝试检测表格
        tables = _detect_tables_in_page(line_infos, img_width, dpi)
        if tables:
            _output_tables(docx, tables, line_infos, img_width, dpi, _join_words_smart)
        else:
            single_type = _classify_paragraph_type_v2(
                single_lines, img_width, img_height, med_body_fs_single,
                med_body_fs_single, avg_word_height, dpi, line_index=0,
            )
            _output_paragraph_v2(docx, single_lines, single_type,
                                 img_width, dpi, _join_words_smart,
                                 med_body_fs=med_body_fs_single)
        return

    # 计算每对相邻行间距
    gaps: list[tuple[int, float]] = []  # (from_line_idx, gap_px)
    for i in range(1, len(line_infos)):
        prev_bottom = line_infos[i - 1][2]
        curr_top = line_infos[i][1]
        gap = curr_top - prev_bottom
        gaps.append((i - 1, max(gap, 0)))

    valid_gaps = [(idx, g) for idx, g in gaps if g > 0]
    if not valid_gaps:
        # 所有行紧贴 → 整个页面当一个大段落
        paragraph_breaks = [0]
    else:
        gap_values = [g for _, g in valid_gaps]
        med_gap = sorted(gap_values)[len(gap_values) // 2]

        # 局部自适应阈值：每个 gap 与前后 3 个 gap 的中位数比较
        LOCAL_WINDOW = 3
        adaptive_thresholds = []
        for i, (idx, g) in enumerate(gaps):
            if g <= 0:
                adaptive_thresholds.append(0)
                continue
            # 取局部窗口内的有效 gap
            local_vals = []
            for j in range(max(0, i - LOCAL_WINDOW), min(len(gaps), i + LOCAL_WINDOW + 1)):
                if gaps[j][1] > 0:
                    local_vals.append(gaps[j][1])
            if not local_vals:
                adaptive_thresholds.append(g)
                continue
            local_med = sorted(local_vals)[len(local_vals) // 2]
            adaptive_thresholds.append(local_med)

        # 段落信号权重系统（0-1 之间，> 0.5 判定为新段落）
        paragraph_breaks = [0]
        for i in range(1, len(line_infos)):
            prev_info = line_infos[i - 1]
            curr_info = line_infos[i]
            gap = curr_info[1] - prev_info[2]
            gap_idx = i - 1

            # 信号 1：间距异常（局部自适应）
            if gap > 0 and gap_idx < len(adaptive_thresholds):
                local_med = adaptive_thresholds[gap_idx]
                if local_med > 0:
                    gap_ratio = gap / local_med
                    # 高于局部中位 1.6× → 强信号
                    gap_signal = min(gap_ratio / 2.0, 1.0)
                else:
                    gap_signal = 1.0 if gap > avg_word_height * 2 else 0.0
            else:
                gap_signal = 0.0

            # 信号 2：对齐突变（居中→左对齐 → 可能是标题后段落）
            prev_align = prev_info[8]
            curr_align = curr_info[8]
            align_signal = 0.0
            if abs(prev_align) < 0.3 and abs(curr_align - (-1.0)) < 0.3:
                # 上一行居中/偏居中 → 当前行左对齐 → 标题后正文
                align_signal = 0.5

            # 信号 3：首行缩进（左边界突然右移超过 1.5 字符宽）
            prev_left = prev_info[5]
            curr_left = curr_info[5]
            indent_signal = 0.0
            char_w = _SCANNED_BODY_FS * dpi / 72  # 固定字符宽度
            indent_px = curr_left - prev_left
            if indent_px > char_w * 1.5:
                # 确认下一行回到正常边界（如果存在）
                if i + 1 < len(line_infos):
                    next_left = line_infos[i + 1][5]
                    if abs(next_left - prev_left) < char_w * 1.0:
                        indent_signal = 0.8
                else:
                    indent_signal = 0.5

            # 综合信号（加权求和，不再依赖不可靠的字号信号）
            total_signal = (
                gap_signal * 0.5 +
                align_signal * 0.25 +
                indent_signal * 0.25
            )

            if total_signal > 0.45:
                paragraph_breaks.append(i)

    # ---- 考试卷边界强制分段：题号、选项、材料、答案等 ----
    # 将已有的段落 break 与考试卷边界行合并，确保每题/选项独立成段
    boundary_breaks = set(paragraph_breaks)
    for idx, li in enumerate(line_infos):
        words = li[0]
        text = _join_words_smart([w[0] for w in words]).strip()
        if _detect_exam_boundary_line(text):
            boundary_breaks.add(idx)
    paragraph_breaks = sorted(boundary_breaks)

    # ==================== 第五步：表格检测 ====================
    # 在分类前检测表格区域，将其从普通段落流中移除
    table_regions = _detect_tables_in_page(line_infos, img_width, dpi)
    table_line_indices: set[int] = set()
    for t in table_regions:
        for li in range(t["start_line"], t["end_line"] + 1):
            table_line_indices.add(li)

    # ==================== 第六步：段落分组 + 分类 ====================

    paragraph_groups = []  # [(start_line, end_line, is_table), ...]
    for p_idx in range(len(paragraph_breaks)):
        start_line = paragraph_breaks[p_idx]
        end_line = (paragraph_breaks[p_idx + 1] - 1
                    if p_idx + 1 < len(paragraph_breaks)
                    else len(line_infos) - 1)

        # 判断该段落是否属于表格区域
        para_lines_in_table = any(li in table_line_indices
                                  for li in range(start_line, end_line + 1))
        paragraph_groups.append((start_line, end_line, para_lines_in_table))

    # 计算页级统计量（扫描件不推测字号，使用固定常量）
    med_body_fs = _SCANNED_BODY_FS
    max_fs_page = _SCANNED_BODY_FS

    # 分类+输出
    for start, end, is_table in paragraph_groups:
        para_lines = line_infos[start:end + 1]
        if not para_lines:
            continue

        if is_table:
            # 表格行 → 找对应表格区域输出
            matched_table = None
            for t in table_regions:
                if t["start_line"] <= start and end <= t["end_line"]:
                    matched_table = t
                    break
            if matched_table:
                # 收集表格内的所有行
                table_all_lines = line_infos[matched_table["start_line"]:
                                             matched_table["end_line"] + 1]
                _output_single_table(docx, matched_table, table_all_lines,
                                     img_width, dpi, _join_words_smart)
                # 跳过同一表格的后续段落
                continue

        # 普通段落分类（使用改进的分类器）
        ptype = _classify_paragraph_type_v2(
            para_lines, img_width, img_height, med_body_fs, max_fs_page,
            avg_word_height, dpi, line_index=start,
        )
        _output_paragraph_v2(
            docx, para_lines, ptype,
            img_width, dpi, _join_words_smart,
            med_body_fs=med_body_fs,
        )


# ═══════════════════════════════════════════════════════════
#  原生 PaddleOCR 行级排版重建（替代伪词级管道）
# ═══════════════════════════════════════════════════════════


def _build_from_ocr_lines(
    docx,
    ocr_lines: list[dict],
    img_width: int,
    img_height: int,
    dpi: int,
):
    """
    原生处理 PaddleOCR 行级结果，重建排版并输出到 docx。

    不再将整行伪装成单个词元，而是直接以行为单位做段落检测。
    核心改进：
    1. OCR 行后处理：修复过度合并（双列粘合）和过度拆分
    2. 硬边界 + 软边界复合段落检测
    3. 语义特征辅助断段（行末标点、起始模式、长度模式）
    """
    import re
    from docx.shared import Pt
    from docx.oxml.ns import qn

    if not ocr_lines:
        return

    # ── 工具函数 ──
    def _is_cjk(ch: str) -> bool:
        cp = ord(ch)
        return (0x4E00 <= cp <= 0x9FFF or 0x3400 <= cp <= 0x4DBF
                or 0x3000 <= cp <= 0x303F or 0xFF00 <= cp <= 0xFFEF
                or 0x3040 <= cp <= 0x309F or 0x30A0 <= cp <= 0x30FF
                or 0xAC00 <= cp <= 0xD7AF)

    def _cjk_ratio(text: str) -> float:
        if not text:
            return 0.0
        return sum(1 for ch in text if _is_cjk(ch)) / len(text)

    def _ends_with_punct(text: str) -> bool:
        """行末是否为自然段结束标点"""
        return text.rstrip() and text.rstrip()[-1] in "。！？…~～）)」》"

    def _starts_with_numbering(text: str) -> bool:
        """行首是否为编号/题号/列表标记"""
        text = text.strip()
        cjk_num = "一二三四五六七八九十"
        patterns = [
            rf"^第[{cjk_num}]+[章节条款]",            # 第一章
            rf"^[{cjk_num}]{{1,2}}[、，]",           # 一、
            r"^\d{1,2}[\.．、]\s*",                    # 1. 2.
            r"^\(\d{1,2}\)\s*",                        # (1)
            r"^[A-Da-d][\.．、]\s*",                     # A.
            r"^[①②③④⑤⑥⑦⑧⑨⑩]",                  # ①
            r"^【(?:答案|解析|注|说明|摘要|关键词)】",
            r"^材料[一二三四五六七八九十]+",
            r"^[（(]\d+[）)]",
            r"^[IVX]+[\.、]",                          # 罗马数字
            r"^[•■□▪▫◆◇●○▶▷→⇒☆★\-·]",            # 项目符号
        ]
        return any(re.match(pat, text) for pat in patterns)

    def _is_exam_boundary(text: str) -> bool:
        """考试卷边界：题号/选项/材料/答案"""
        text = text.strip()
        patterns = [
            r"^\d{1,2}[\.．]\s*",
            r"^\(\d{1,2}\)\s*",
            r"^【答案】", r"^【解析】",
            r"^材料[一二三四五六七八九十]+\s*",
            r"^[A-Da-d][\.．、]\s*",
            r"^\(\d+分\)",
        ]
        return any(re.match(pat, text) for pat in patterns)

    # ═══════════════════════════════════════════════
    #  步骤1：OCR 后处理 — 修复行切分错误
    # ═══════════════════════════════════════════════
    fixed_lines = _postprocess_ocr_lines(ocr_lines, img_width)

    if not fixed_lines:
        return

    # ═══════════════════════════════════════════════
    #  步骤2：计算全局统计量
    # ═══════════════════════════════════════════════
    line_heights = [ol["height"] for ol in fixed_lines]
    med_height = sorted(line_heights)[len(line_heights) // 2] if line_heights else 20

    # 行间距
    gaps = []
    for i in range(1, len(fixed_lines)):
        prev_bottom = fixed_lines[i - 1]["bottom"]
        curr_top = fixed_lines[i]["top"]
        gaps.append(max(curr_top - prev_bottom, 0))
    med_gap = sorted(gaps)[len(gaps) // 2] if gaps else med_height

    # 正文行高（用于缩进计算）
    char_w_px = _SCANNED_BODY_FS * dpi / 72  # 一个字符的像素宽度

    # ═══════════════════════════════════════════════
    #  步骤3：段落边界检测（硬边界 + 软边界）
    # ═══════════════════════════════════════════════

    # 3a. 硬边界：必定分段
    hard_breaks: set[int] = set()
    hard_breaks.add(0)  # 首行总是段落开始

    # 大间距：超过中位间距 3 倍
    big_gap_threshold = max(med_gap * 3, med_height * 2.5)
    for i in range(1, len(fixed_lines)):
        prev_bottom = fixed_lines[i - 1]["bottom"]
        curr_top = fixed_lines[i]["top"]
        gap = curr_top - prev_bottom
        if gap > big_gap_threshold:
            hard_breaks.add(i)

    # 考试卷边界
    for i, ol in enumerate(fixed_lines):
        if _is_exam_boundary(ol["text"]):
            hard_breaks.add(i)

    # 编号起始：任何编号行都是新段落的开始
    # 例外：如果前后行都是同一类型的紧凑短编号（如"A. B. C. 选项"），不强制拆分
    for i in range(1, len(fixed_lines)):
        if _starts_with_numbering(fixed_lines[i]["text"]):
            prev_text = fixed_lines[i - 1]["text"].strip()
            curr_text = fixed_lines[i]["text"].strip()
            # 如果当前是短选项（A./B.等）且前一行也是同类短选项，则属于同一选择题的行内选项
            is_short_option = bool(re.match(r"^[A-Da-d][\.．、]", curr_text))
            prev_is_option = bool(re.match(r"^[A-Da-d][\.．、]", prev_text))
            if is_short_option and prev_is_option and len(curr_text) < 20 and len(prev_text) < 20:
                continue  # 同类短选项，不强制分段（可能本就应在同一选择题内）
            hard_breaks.add(i)

    # 行末标点结尾 + 下一行缩进 → 可能是新段
    for i in range(1, len(fixed_lines)):
        prev_text = fixed_lines[i - 1]["text"].strip()
        curr_left = fixed_lines[i]["left"]
        prev_left = fixed_lines[i - 1]["left"]
        if _ends_with_punct(prev_text) and (curr_left - prev_left) > char_w_px * 1.2:
            hard_breaks.add(i)

    # 3b. 软边界：综合信号判定（仅对非硬边界位置）
    all_breaks = set(hard_breaks)
    for i in range(1, len(fixed_lines)):
        if i in hard_breaks:
            continue

        prev = fixed_lines[i - 1]
        curr = fixed_lines[i]
        gap = curr["top"] - prev["bottom"]

        # 信号 1：间距信号 — 相对于局部中位数
        # 取前后各 2 行的间距中位数
        local_gaps = []
        for j in range(max(1, i - 2), min(len(fixed_lines) - 1, i + 2)):
            if j < len(fixed_lines) - 1:
                lg = fixed_lines[j + 1]["top"] - fixed_lines[j]["bottom"]
                if lg > 0:
                    local_gaps.append(lg)
        local_med = sorted(local_gaps)[len(local_gaps) // 2] if local_gaps else med_gap
        gap_signal = min(gap / max(local_med, 1) / 2.0, 1.0) if local_med > 0 else 0.0

        # 信号 2：行末标点信号
        punct_signal = 0.6 if _ends_with_punct(prev["text"]) else 0.0

        # 信号 3：下一行缩进信号
        indent_px = curr["left"] - prev["left"]
        indent_signal = 0.5 if indent_px > char_w_px * 1.2 else 0.0

        # 信号 4：长度突变信号（短行→长行，可能标题→正文）
        prev_len = len(prev["text"].strip())
        curr_len = len(curr["text"].strip())
        length_signal = 0.0
        if prev_len > 0 and curr_len > prev_len * 2.5:
            length_signal = 0.4

        # 信号 5：对齐变化（居中→左对齐）
        prev_center = abs((prev["left"] + prev["right"]) / 2 - img_width / 2)
        curr_center = abs((curr["left"] + curr["right"]) / 2 - img_width / 2)
        align_signal = 0.3 if (prev_center < img_width * 0.1
                               and curr_center > img_width * 0.15) else 0.0

        # 加权综合
        total_signal = (
            gap_signal * 0.35
            + punct_signal * 0.25
            + indent_signal * 0.15
            + length_signal * 0.15
            + align_signal * 0.10
        )
        if total_signal > 0.40:
            all_breaks.add(i)

    paragraph_starts = sorted(all_breaks)

    # ═══════════════════════════════════════════════
    #  步骤4：组装段落并输出
    # ═══════════════════════════════════════════════
    for p_idx in range(len(paragraph_starts)):
        start = paragraph_starts[p_idx]
        end = (paragraph_starts[p_idx + 1] - 1
               if p_idx + 1 < len(paragraph_starts)
               else len(fixed_lines) - 1)

        para_lines = fixed_lines[start:end + 1]
        if not para_lines:
            continue

        _output_ocr_paragraph(docx, para_lines, img_width, char_w_px)


# ── OCR 行后处理：修复切分错误 ──

def _postprocess_ocr_lines(
    ocr_lines: list[dict], img_width: int
) -> list[dict]:
    """
    修复 PaddleOCR 常见的行切分错误：

    1. 过度合并：同一行检测框内包含多列文字（宽框内有 >30% 页宽的水平空白）
       → 按空白间隙拆分为多行
    2. 过度拆分：连续两行左对齐、行文本没有句末标点、间距正常
       → 合并为一行
    3. 修复后的行按 y 坐标重新排序
    """
    import re

    if not ocr_lines:
        return []

    # 先按 top 排序，确保顺序处理
    sorted_lines = sorted(ocr_lines, key=lambda ol: (ol["top"], ol["left"]))

    result: list[dict] = []

    # ── 第1遍：检测过度合并的行（宽框内有多列空白）──
    for ol in sorted_lines:
        text = ol["text"]
        width = ol["width"]
        # 宽度 > 60% 页宽 → 可能合并了多列
        if width > img_width * 0.55 and len(text) > 15:
            # 在空格的宽间隙处拆分（>6个字符宽度的空白）
            parts = _split_merged_line_by_gaps(text, width, img_width)
            if len(parts) > 1:
                # 按比例分配坐标
                part_width = width / len(parts)
                for pi, part in enumerate(parts):
                    result.append({
                        "text": part.strip(),
                        "left": ol["left"] + int(pi * part_width * 0.95),
                        "top": ol["top"],
                        "right": ol["left"] + int((pi + 1) * part_width * 0.95),
                        "bottom": ol["bottom"],
                        "width": int(part_width * 0.9),
                        "height": ol["height"],
                        "conf": ol.get("conf", 0.9),
                        "line_idx": ol.get("line_idx", 0),
                    })
                continue
        result.append(ol)

    # ── 第2遍：检测过度拆分的行（连续短行本应属于同一行）──
    # 条件：相邻两行 左边界接近 + 间距 < 1.2x 行高 + 前一行不以标点结尾
    merged: list[dict] = []
    i = 0
    while i < len(result):
        current = dict(result[i])
        # 尝试与后续行合并
        while i + 1 < len(result):
            next_ol = result[i + 1]
            gap = next_ol["top"] - current["bottom"]
            left_diff = abs(current["left"] - next_ol["left"])
            same_line_gap = current["height"] * 1.1

            # 合并条件：间距紧 + 左对齐 + 不是自然段尾 + 不是考试边界 + 不是表格式数据
            curr_tokens = len(current["text"].split())
            next_tokens = len(next_ol["text"].split())
            is_tabular = curr_tokens >= 3 or next_tokens >= 3

            if (gap < same_line_gap
                    and left_diff < current["height"] * 1.5
                    and not (current["text"].rstrip()
                             and current["text"].rstrip()[-1] in "。！？…")
                    and not _is_line_boundary(current["text"])
                    and not _is_line_boundary(next_ol["text"])
                    and not is_tabular):
                # 合并文本（CJK 不加空格，Latin 加空格）
                curr_text = current["text"].strip()
                next_text = next_ol["text"].strip()
                if (curr_text and next_text
                        and curr_text[-1].isascii()
                        and next_text[0].isascii()):
                    current["text"] = curr_text + " " + next_text
                else:
                    current["text"] = curr_text + next_text
                current["right"] = max(current["right"], next_ol["right"])
                current["bottom"] = next_ol["bottom"]
                current["width"] = current["right"] - current["left"]
                current["height"] = current["bottom"] - current["top"]
                i += 1  # 跳过被合并的行
            else:
                break
        merged.append(current)
        i += 1

    # ── 第3遍：修复行内间隙过大的单行（可能仍残留合并）──
    final: list[dict] = []
    for ol in merged:
        text = ol["text"]
        # 检查是否有中文字符之间的超大空白（通常是非等宽字体造成的空格折叠）
        text = re.sub(r'(?<=[\u4e00-\u9fff])\s{2,}(?=[\u4e00-\u9fff])', '', text)
        # 检查是否有多余的前后空格
        text = text.strip()
        if text:
            ol["text"] = text
            final.append(ol)

    return final


def _split_merged_line_by_gaps(
    text: str, line_width: int, img_width: int
) -> list[str]:
    """
    对宽行按文本内间隙拆分。利用 OCR 文本中空格的分布推断列边界。
    如果存在 ≥ 3 个连续空格且位置相对均匀 → 判定为多列合并。
    """
    import re

    # 找连续 ≥ 3 个空格的位置
    gap_positions = []
    for m in re.finditer(r'\s{3,}', text):
        gap_positions.append((m.start(), m.end()))

    if not gap_positions:
        return [text]

    # 过滤：只保留跨距相对均匀的间隙
    if len(gap_positions) == 1:
        # 只有一个大间隙 → 可能分为 2 列
        s, e = gap_positions[0]
        left_part = text[:s].strip()
        right_part = text[e:].strip()
        if left_part and right_part and len(left_part) >= 2 and len(right_part) >= 2:
            return [left_part, right_part]
        return [text]

    # 多个间隙 → 检查是否均匀分布
    positions = [(g[0] + g[1]) // 2 for g in gap_positions]
    positions.append(len(text))  # 文本末尾作为最后一个边界
    segments = []
    prev = 0
    for pos in positions:
        seg = text[prev:pos].strip()
        if seg:
            segments.append(seg)
        prev = pos

    # ── 严格的拆分验证 ──
    if len(segments) < 2:
        return [text]

    seg_lens = [len(s) for s in segments]
    avg_seg_len = sum(seg_lens) / len(seg_lens)
    short_count = sum(1 for sl in seg_lens if sl < 5)

    # 拒绝条件 1：大多数段太短 → 是表头文字，不是合并列
    if short_count > len(segments) / 2:
        return [text]

    # 拒绝条件 2：平均段长 < 6 → 段内容太少，不适合拆分
    if avg_seg_len < 6:
        return [text]

    # 拒绝条件 3：拆分太多（> 5 列）→ 很可能是表头误拆
    if len(segments) > 5:
        return [text]

    # 拒绝条件 4：段长悬殊过大（最大/最小 ≥ 4）
    if max(seg_lens) / max(min(seg_lens), 1) >= 4:
        return [text]

    return segments


# ── 输出 OCR 段落到 docx ──

def _output_ocr_paragraph(
    docx,
    para_lines: list[dict],
    img_width: int,
    char_w_px: float,
):
    """将已分组为段落的 OCR 行输出为 Word 段落。"""
    from docx.shared import Pt
    import re

    if not para_lines:
        return

    # 合并段内所有行的文本
    line_texts = [ol["text"].strip() for ol in para_lines]
    line_texts = [t for t in line_texts if t]
    if not line_texts:
        return

    # 检查是否为列表/编号项
    first_text = line_texts[0]
    is_list = _is_list_item(first_text) or _starts_with_numbering(first_text)

    # 检查是否为考试卷选项
    is_exam_option = bool(re.match(r"^[A-Da-d][\.．、]", first_text))

    # 检查 CJK 占比
    total_text = "".join(line_texts)
    cjk_chars = sum(1 for ch in total_text if '\u4e00' <= ch <= '\u9fff'
                    or '\u3400' <= ch <= '\u4dbf')
    is_cjk_dominant = len(total_text) > 0 and cjk_chars / len(total_text) > 0.40

    # 检查是否居中（可能是标题）
    avg_center_offset = sum(
        abs((ol["left"] + ol["right"]) / 2 - img_width / 2)
        for ol in para_lines
    ) / len(para_lines)
    is_centered = avg_center_offset < img_width * 0.12

    # 创建段落
    p = docx.add_paragraph()

    # 设置段落格式
    if is_list or is_exam_option:
        # 列表项：悬挂缩进
        p.paragraph_format.left_indent = Pt(_SCANNED_BODY_FS)
        p.paragraph_format.first_line_indent = Pt(-_SCANNED_BODY_FS)
    elif is_centered and len(total_text) < 60:
        # 短居中文本 → 标题效果（无缩进，居中）
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    elif is_cjk_dominant and len(total_text) > 20:
        # CJK 长段落：首行缩进 2 字符
        p.paragraph_format.first_line_indent = Pt(_SCANNED_BODY_FS * 2)

    # 输出文本：拆分行内多题号/多选项
    for li, text in enumerate(line_texts):
        # CJK 段落拆题号和选项
        if is_cjk_dominant:
            parts = _split_multi_questions_in_line(text)
            sub_parts = []
            for part in parts:
                sub_parts.extend(_split_exam_options_in_line(part))
        else:
            sub_parts = [text]

        for pi, part in enumerate(sub_parts):
            if not part.strip():
                continue
            run = p.add_run(part)
            run.font.size = Pt(_SCANNED_BODY_FS)
            # 行内选项之间换行，不同 OCR 行之间换行
            if pi < len(sub_parts) - 1 or li < len(line_texts) - 1:
                p.add_run("\n")

    # 确保段落至少有一个换行（不为空）
    if not p.runs:
        p.add_run("")


def _starts_with_numbering(text: str) -> bool:
    """行首是否为编号/题号/列表标记（模块级版本）"""
    import re
    text = text.strip()
    cjk_num = "一二三四五六七八九十"
    patterns = [
        rf"^第[{cjk_num}]+[章节条款]",
        rf"^[{cjk_num}]{{1,2}}[、，]",
        r"^\d{1,2}[\.．、]\s*",
        r"^\(\d{1,2}\)\s*",
        r"^[A-Da-d][\.．、]\s*",
        r"^[①②③④⑤⑥⑦⑧⑨⑩]",
        r"^【(?:答案|解析|注|说明|摘要|关键词)】",
        r"^材料[一二三四五六七八九十]+",
        r"^[（(]\d+[）)]",
        r"^[IVX]+[\.、]",
        r"^[•■□▪▫◆◇●○▶▷→⇒☆★\-·]",
    ]
    return any(re.match(pat, text) for pat in patterns)


def _is_line_boundary(text: str) -> bool:
    """检测行是否为不可合并的边界行（题号、选项、答案行等）"""
    import re
    text = text.strip()
    patterns = [
        r"^\d{1,2}[\.．]\s*",           # 24. 25.
        r"^\(\d{1,2}\)\s*",              # (1) (2)
        r"^【答案】",                      # 【答案】
        r"^【解析】",                      # 【解析】
        r"^材料[一二三四五六七八九十]+\s*", # 材料一
        r"^[A-Da-d][\.．、]\s*",           # A. B.
    ]
    return any(re.match(pat, text) for pat in patterns)


# ──────────────────────────────────────────────
#  段落分类器 v2（位置权重 + 字号分布 + 多特征融合）
# ──────────────────────────────────────────────

def _classify_paragraph_type_v2(
    para_lines: list,
    img_width: int,
    img_height: int,
    med_body_fs: float,
    max_fs_page: float,
    med_line_gap: float,
    dpi: int,
    line_index: int = 0,
) -> str:
    """
    段落分类器（已简化）。

    不做强制排版处理后，所有段落统一按纯文本输出，
    不再区分标题/列表/正文。保留此函数仅为了调用兼容。
    """
    return "body"


def _detect_heading_numbering(text: str) -> bool:
    """检测文本是否包含标题编号模式（第X章、一、...、1.1 等）"""
    import re
    text = text.strip()
    cjk_num = "一二三四五六七八九十"
    patterns = [
        rf"^第[{cjk_num}]+[章节条]",
        rf"^[{cjk_num}]{{1,2}}[、，]",
        r"^\d+(\.\d+)*[\s、\t]+",
        r"^[（(][\d]+[）)]",
        r"^[①②③④⑤⑥⑦⑧⑨⑩]",
        r"^[IVX]+[\.、\s]",
    ]
    for pat in patterns:
        if re.match(pat, text):
            return True
    return False


def _is_list_item_v2(
    text: str, first_left: float, img_width: int,
    med_body_fs: float, dpi: int,
) -> bool:
    """检测段落是否为列表项（编号/项目符号开头 + 缩进辅助判断）"""
    import re

    text = text.strip()
    if not text:
        return False

    cjk_number = (
        "一|二|三|四|五|六|七|八|九|十|"
        "十一|十二|十三|十四|十五|十六|十七|十八|十九|二十"
    )

    numbered_patterns = [
        rf"^第[{cjk_number}]+[章条节款]",
        rf"^[{cjk_number}]、",
        rf"^[（(][{cjk_number}]+[）)]",
        r"^\d+[\.\、\)）]",
        r"^\d+\.\d+[\s\、]?",     # 多级编号
        r"^[①②③④⑤⑥⑦⑧⑨⑩]",
        r"^[（(]\d+[）)]",
        r"^[IVX]+[\.\、]",
        r"^[a-zA-Z][\.\、\)]",   # a) b) 等
    ]
    for pat in numbered_patterns:
        if re.match(pat, text):
            return True

    bullet_chars = ('•', '■', '□', '▪', '▫', '◆', '◇', '●', '○',
                    '▶', '▷', '→', '⇒', '☆', '★', '-', '·')
    if text[0] in bullet_chars and len(text) > 2:
        return True

    # 额外：悬挂缩进检测（正文左缩进+无编号可能是列表续行，不回退）
    char_w = med_body_fs * dpi / 72
    if first_left > char_w * 2.5 and first_left < img_width * 0.4:
        # 明显左缩进 + 文本短 → 可能是列表
        if len(text) < 80:
            return True

    return False


# ──────────────────────────────────────────────
#  表格检测（x/y 网格对齐分析）
# ──────────────────────────────────────────────

def _detect_tables_in_page(
    line_infos: list, img_width: int, dpi: int,
) -> list[dict]:
    """
    在页面中检测表格区域（改进版）。

    核心思路：
    表格行由多个文本单元格组成，单元格之间有大片垂直空白。
    如果连续多行在这些空白位置对齐，则判定为表格。

    算法：
    1. 对每行按 x 排序，找出相邻词之间的"大间隙"
    2. 比较相邻行的间隙位置，对齐则视为同一表格行
    3. 连续 ≥3 行且共享 ≥1 个对齐间隙 → 表格区域
    4. 根据对齐间隙划分列，计算每列的词中心作为列中心
    """
    if len(line_infos) < 3:
        return []

    def _row_gaps(words_for_line: list) -> list[float]:
        """返回一行中可作为列分隔的间隙中点列表。"""
        if len(words_for_line) < 2:
            return []
        sorted_words = sorted(words_for_line, key=lambda w: w[1])

        widths = [w[3] for w in sorted_words]
        avg_width = sum(widths) / len(widths) if widths else 0
        # 间隙阈值：适度宽松，以捕获印刷体表格中的中等列间距
        gap_threshold = max(avg_width * 0.7, img_width * 0.015)

        gaps = []
        for i in range(len(sorted_words) - 1):
            right_i = sorted_words[i][1] + sorted_words[i][3]
            left_j = sorted_words[i + 1][1]
            gap = left_j - right_i
            if gap > gap_threshold:
                gaps.append((right_i + left_j) / 2)
        return gaps

    def _gaps_match(gaps_a: list[float], gaps_b: list[float], tolerance: float) -> bool:
        """判断两组间隙是否有至少一对对齐（允许 tolerance）。"""
        if not gaps_a or not gaps_b:
            return False
        for a in gaps_a:
            for b in gaps_b:
                if abs(a - b) <= tolerance:
                    return True
        return False

    def _find_common_gaps(gaps_list: list[list[float]], tolerance: float) -> list[float]:
        """找多行共同稳定出现的间隙位置（投票法）。"""
        all_gaps = []
        for gaps in gaps_list:
            all_gaps.extend(gaps)
        if not all_gaps:
            return []

        all_gaps.sort()
        clusters = [[all_gaps[0]]]
        for x in all_gaps[1:]:
            if abs(x - clusters[-1][-1]) <= tolerance:
                clusters[-1].append(x)
            else:
                clusters.append([x])

        # 只保留出现次数 >= 行数 60% 的簇（比 50% 更严格，减少误检）
        min_count = max(2, len(gaps_list) * 0.6)
        common = [sum(c) / len(c) for c in clusters if len(c) >= min_count]
        return sorted(common)

    def _col_centers_from_gaps(
        table_lines: list, gap_positions: list[float], img_w: int,
    ) -> list[float]:
        """根据间隙边界划分列，返回每列的词平均 x 中心。"""
        boundaries = sorted([0] + gap_positions + [img_w])
        col_words: dict[int, list[float]] = {i: [] for i in range(len(boundaries) - 1)}

        for line in table_lines:
            for w in line[0]:
                w_left = w[1]
                w_right = w_left + w[3]
                w_center = (w_left + w_right) / 2
                for ci in range(len(boundaries) - 1):
                    if boundaries[ci] <= w_center < boundaries[ci + 1]:
                        col_words[ci].append(w_center)
                        break

        centers = []
        for ci in sorted(col_words.keys()):
            vals = col_words[ci]
            if vals:
                centers.append(sum(vals) / len(vals))
        return centers

    TOLERANCE = img_width * 0.03

    # 1. 计算每行的大间隙
    rows_gaps: list[tuple[int, list[float]]] = []
    for li_idx, line_info in enumerate(line_infos):
        words = line_info[0]
        rows_gaps.append((li_idx, _row_gaps(words)))

    # 2. 扫描连续行，找表格区域
    table_regions: list[dict] = []
    in_table = False
    table_start = 0
    table_gaps_list: list[list[float]] = []

    for i, (li_idx, gaps) in enumerate(rows_gaps):
        if not in_table:
            if len(gaps) >= 1:
                in_table = True
                table_start = li_idx
                table_gaps_list = [gaps]
        else:
            prev_gaps = table_gaps_list[-1]
            if _gaps_match(prev_gaps, gaps, TOLERANCE):
                table_gaps_list.append(gaps)
            else:
                # 表格结束
                if len(table_gaps_list) >= 3:
                    common_gaps = _find_common_gaps(table_gaps_list, TOLERANCE)
                    if len(common_gaps) >= 1:
                        table_lines = line_infos[table_start:li_idx]
                        col_centers = _col_centers_from_gaps(
                            table_lines, common_gaps, img_width,
                        )
                        if len(col_centers) >= 2:
                            table_regions.append({
                                "start_line": table_start,
                                "end_line": li_idx - 1,
                                "columns": col_centers,
                                "num_cols": len(col_centers),
                                "rows": len(table_gaps_list),
                            })
                in_table = False
                if len(gaps) >= 1:
                    table_start = li_idx
                    table_gaps_list = [gaps]

    # 处理尾部表格
    if in_table and len(table_gaps_list) >= 3:
        common_gaps = _find_common_gaps(table_gaps_list, TOLERANCE)
        if len(common_gaps) >= 1:
            table_lines = line_infos[table_start:len(line_infos)]
            col_centers = _col_centers_from_gaps(table_lines, common_gaps, img_width)
            if len(col_centers) >= 2:
                table_regions.append({
                    "start_line": table_start,
                    "end_line": len(line_infos) - 1,
                    "columns": col_centers,
                    "num_cols": len(col_centers),
                    "rows": len(table_gaps_list),
                })

    # 合并重叠或紧邻区域
    if not table_regions:
        return []
    merged = [table_regions[0]]
    for t in table_regions[1:]:
        prev = merged[-1]
        if t["start_line"] <= prev["end_line"] + 1:
            prev["end_line"] = max(prev["end_line"], t["end_line"])
            prev["rows"] = max(prev["rows"], t["rows"])
            prev["num_cols"] = max(prev["num_cols"], t["num_cols"])
        else:
            merged.append(t)

    # 斜线表头扩展：若表格上方有"类别/项目"等跨列表头文字，纳入表格
    extended = []
    for t in merged:
        start = t["start_line"]
        if start > 0:
            prev_line = line_infos[start - 1]
            prev_text = "".join(w[0] for w in prev_line[0]).strip()
            if ("类别" in prev_text or "项目" in prev_text) and len(prev_line[0]) <= 3:
                start = start - 1
                t = {
                    "start_line": start,
                    "end_line": t["end_line"],
                    "columns": t["columns"],
                    "num_cols": t["num_cols"],
                    "rows": t["end_line"] - start + 1,
                }
        extended.append(t)

    return extended



def _output_single_table(
    docx,
    table_info: dict,
    all_lines: list,
    img_width: int,
    dpi: int,
    join_func,
):
    """将检测到的单个表格区域输出为 DOCX 原生表格"""
    from docx.shared import Pt, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    start = table_info["start_line"]
    end = table_info["end_line"]
    num_cols = table_info["num_cols"]
    col_centers = table_info["columns"]

    table_lines = all_lines  # 已经是表格范围内的行

    if len(table_lines) < 2:
        # 行不够 → 回退到普通段落
        for li in table_lines:
            p = docx.add_paragraph()
            words_text = [w[0] for w in li[0]]
            p.add_run(join_func(words_text))
        return

    # ── 为每行的每个词分配列索引 ──
    def _assign_cell(words_for_line, col_centers, img_w):
        """返回 dict: {col_idx: [word_texts]}"""
        cells: dict[int, list[str]] = {}
        for w in words_for_line:
            w_left = w[1]
            best_col = 0
            best_dist = img_w
            for ci, cx in enumerate(col_centers):
                dist = abs(w_left - cx)
                if dist < best_dist:
                    best_dist = dist
                    best_col = ci
            if best_dist < img_w * 0.06:
                cells.setdefault(best_col, []).append(w[0])
        return cells

    # 生成二维网格
    grid: list[dict[int, str]] = []
    for li in table_lines:
        words = li[0]
        cell_dict = _assign_cell(words, col_centers, img_width)
        row_data = {}
        for ci in range(num_cols):
            texts = cell_dict.get(ci, [])
            row_data[ci] = join_func(texts) if texts else ""
        grid.append(row_data)

    # 过滤全空行和全空列
    non_empty_rows = [r for r in grid if any(v.strip() for v in r.values())]
    if not non_empty_rows:
        return

    # 找非空列
    active_cols = sorted(set(
        ci for r in non_empty_rows
        for ci, txt in r.items() if txt.strip()
    ))
    if len(active_cols) < 2:
        # 只一列有效 → 普通段落
        for r in non_empty_rows:
            p = docx.add_paragraph()
            texts = [r.get(ci, "") for ci in active_cols]
            p.add_run(" ".join(texts))
        return

    # 重新映射列索引
    col_map = {old: new for new, old in enumerate(active_cols)}
    actual_cols = len(active_cols)

    # 检测并处理斜线表头：第一行仅含"类别/项目"等跨列表头文字
    # 斜线表头在 Word 中用"合并首行所有单元格"来近似还原
    diagonal_header_text = ""
    header_row_offset = 0
    if non_empty_rows:
        first_row = non_empty_rows[0]
        first_text = "".join(first_row.get(ci, "") for ci in active_cols).strip()
        if (
            ("类别" in first_text or "项目" in first_text)
            and len(first_text) <= 6
            and len(non_empty_rows) > 1
        ):
            diagonal_header_text = first_text
            non_empty_rows = non_empty_rows[1:]
            header_row_offset = 1

    if not non_empty_rows:
        return

    # 创建 DOCX 表格
    table = docx.add_table(
        rows=len(non_empty_rows) + header_row_offset, cols=actual_cols
    )
    table.style = 'Table Grid'

    # 输出斜线表头（合并首行）
    if header_row_offset and diagonal_header_text:
        for ci in range(1, actual_cols):
            table.cell(0, 0).merge(table.cell(0, ci))
        cell = table.cell(0, 0)
        cell.text = ""
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(diagonal_header_text)
        run.font.bold = True
        run.font.size = Pt(9)
        run.font.name = 'SimSun'
        _set_east_asian_font(run, 'SimSun')

    for ri, row_data in enumerate(non_empty_rows):
        table_ri = ri + header_row_offset
        for ci_old, ci_new in col_map.items():
            cell_text = row_data.get(ci_old, "")
            cell = table.cell(table_ri, ci_new)
            cell.text = ""
            p = cell.paragraphs[0]
            run = p.add_run(cell_text)
            # 表头行（原第一行数据）加粗体
            if ri == 0:
                run.font.bold = True
                run.font.size = Pt(9)
            else:
                run.font.size = Pt(9)
            run.font.name = 'SimSun'
            _set_east_asian_font(run, 'SimSun')


    # 表格后添加空行分隔
    p_sep = docx.add_paragraph()
    p_sep.paragraph_format.space_before = Pt(4)
    p_sep.paragraph_format.space_after = Pt(4)


def _output_tables(
    docx,
    tables: list[dict],
    line_infos: list,
    img_width: int,
    dpi: int,
    join_func,
):
    """输出所有检测到的表格"""
    for t in tables:
        table_lines = line_infos[t["start_line"]:t["end_line"] + 1]
        _output_single_table(docx, t, table_lines, img_width, dpi, join_func)

# ──────────────────────────────────────────────
#  段落输出 v2（支持标题样式 + 正文排版）
#  注：表格仅由词级检测路径处理，OCR 路径直接输出文字。
# ──────────────────────────────────────────────

def _output_paragraph_v2(
    docx,
    line_infos: list,
    para_type: str,
    img_width: int,
    dpi: int,
    join_func,
    paragraph_spacing_pt: float = 0,
    med_body_fs: float = 10.5,
):
    """
    将多行输出为一个 Word 段落（纯文本，无强制排版）。

    只关注换行和空格，不做任何对齐/间距/缩进/标题样式处理。
    LineInfo = tuple[words, top, bottom, height, fs_pt, left, right, bold, align_score]
    """
    from docx.shared import Pt

    if not line_infos:
        return

    # ── 判断是否为 CJK 主导段落 ──
    line_texts = []
    for words, _, _, _, _, _, _, _, _ in line_infos:
        word_texts = [w[0] for w in words]
        lt = join_func(word_texts).strip()
        if lt:
            line_texts.append(lt)

    total_chars = sum(len(t) for t in line_texts)
    cjk_chars = sum(
        1 for t in line_texts for ch in t
        if ('\u4e00' <= ch <= '\u9fff' or '\u3400' <= ch <= '\u4dbf')
    )
    is_cjk_dominant = total_chars > 0 and (cjk_chars / total_chars) > 0.5

    # ── 判断段落类型：列表项 / 中文正文 / 其他 ──
    is_list_like = False
    if line_texts:
        is_list_like = _is_list_item(line_texts[0])

    # ── 单一纯文本段落：无对齐、无间距 ──
    # 对列表项使用悬挂缩进：编号在左边界，换行文字缩进 1 字符
    # 对中文正文使用首行缩进 2 字符
    p = docx.add_paragraph()
    if is_list_like:
        p.paragraph_format.left_indent = Pt(med_body_fs)
        p.paragraph_format.first_line_indent = Pt(-med_body_fs)
    elif is_cjk_dominant:
        p.paragraph_format.first_line_indent = Pt(med_body_fs * 2)

    for line_idx, (words, _, _, _, _, _, _, _, _) in enumerate(line_infos):
        word_texts = [w[0] for w in words]
        line_text = join_func(word_texts).strip()
        if not line_text:
            continue

        # CJK 段落行内可能有多题号或多选项，拆成多行
        if is_cjk_dominant:
            q_parts = _split_multi_questions_in_line(line_text)
            sub_parts = []
            for q_part in q_parts:
                sub_parts.extend(_split_exam_options_in_line(q_part))
        else:
            sub_parts = [line_text]

        for part_idx, part in enumerate(sub_parts):
            if not part.strip():
                continue
            run = p.add_run(part)
            # 行内多选项之间换行；不同 OCR 行之间换行
            if part_idx < len(sub_parts) - 1 or line_idx < len(line_infos) - 1:
                p.add_run("\n")


def _set_east_asian_font(run, font_name: str):
    """设置 run 的中文/东亚字体"""
    from docx.oxml.ns import qn
    try:
        rpr = run._element.get_or_add_rPr()
        rFonts = rpr.find(qn('w:rFonts'))
        if rFonts is None:
            from lxml import etree
            rFonts = etree.SubElement(rpr, qn('w:rFonts'))
        rFonts.set(qn('w:eastAsia'), font_name)
    except Exception:
        pass





def get_pdf_page_count(file_path: str) -> int:
    """获取 PDF 页数（用于预览）"""
    try:
        import fitz
        doc = fitz.open(file_path)
        count = doc.page_count
        doc.close()
        return count
    except Exception:
        return 0
